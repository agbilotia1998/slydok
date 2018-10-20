from docx import Document
import pprint
from collections import OrderedDict
import io
import xml.etree.ElementTree as ET
from images import extract_images
import re

pp = pprint.PrettyPrinter(indent=2)

filename = 'c.docx'
document = Document(filename)
images = extract_images(filename)
image_count = 1

def convert(p):
    return p.text

def find_all_font_sizes(paragraphs):
    sizes = set()

    for p in paragraphs:
        font_size = p.style.font.size
        sizes.add(font_size)

    return sorted(list(sizes), reverse=True)

def analyze_by_size(paragraphs, sizes):
    cur_size = sizes[0]
    lines = []
    headings = 0
    prev_heading = ''
    d= OrderedDict()
    d["children"] = {}
    d["images"] = []
    d["ul"] = []
    d["text"] = []

    image_count = 1

    for p in paragraphs:
        if len(p._p.r_lst[0].drawing_lst):
            print p._p.r_lst[0].xml, dir(p._p.r_lst[0]), p._p.r_lst[0].drawing_lst
        xml_string = str(p._p.xml)
        result = re.search('graphic', xml_string)

        if result is not None:
            image_name = 'image{0}'.format(image_count)
            print(image_name)

            if 'images' in d:
                d['images'].append(image_name)
            else:
                d['images'] = [image_name]

            image_count += 1

        if cur_size is not None and p.style.font.size == cur_size:
            if headings > 0:
                d['children'][prev_heading] = analyze_by_size(lines, sizes[1:])
            elif headings == 0 and len(lines) > 0:
                d['text'] = map(convert, lines)

            prev_heading = p.text
            d['children'][prev_heading] = {}
            headings += 1
            lines = []
        else:
            lines.append(p)

    if headings > 0:
        d['children'][prev_heading] = analyze_by_size(lines, sizes[1:])
    else:
        d['text'] = map(convert, lines)

    return d

def analyze_by_heading(paragraphs, level):
    heading = 'Heading {0}'.format(level)
    lines = []
    headings = 0
    prev_heading = ''
    d = OrderedDict()

    for p in paragraphs:
        if p.style.name == heading:
            if headings > 0:
                d[prev_heading] = analyze_by_heading(lines, level + 1)
            elif headings == 0 and len(lines) > 0:
                d['intro_text'] = map(convert, lines)

            prev_heading = p.text
            d[prev_heading] = {}
            headings += 1
            lines = []
        else:
            lines.append(p)

    if headings > 0:
        d[prev_heading] = analyze_by_heading(lines, level + 1)

        return d
    else:
        return map(convert, lines)


def remove_extra_images(store):
    if 'images' in store:
        for image in store['images']:
            for k in store.keys():
                if isinstance(store[k], dict) and 'images' in store[k] and image in store[k]['images']:
                    store['images'].remove(image)
# store = analyze(document.paragraphs, 1)
# pp.pprint(store)

title = document.paragraphs[0].text
l = find_all_font_sizes(document.paragraphs[1:])
store = analyze_by_size(document.paragraphs[1:], l)
remove_extra_images(store)
pp.pprint(store)
